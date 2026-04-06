from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import datetime

doc = Document()

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0, 51, 102)
    hs.font.name = 'Times New Roman'
    hs.font.bold = True
    if i == 1: hs.font.size = Pt(18)
    elif i == 2: hs.font.size = Pt(14)
    else: hs.font.size = Pt(12)

def add_centered(text, size=12, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.name = 'Times New Roman'
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_para(text, bold=False, italic=False, align=None, size=12, space_after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(space_after)
    return p

# ========== PAGE DE GARDE ==========
for _ in range(4): doc.add_paragraph()
add_centered("RAPPORT DE STAGE DE FIN D'ÉTUDES", 22, True, RGBColor(0, 51, 102))
doc.add_paragraph()
add_centered("Présenté pour l'obtention du diplôme", 13)
add_centered("Filière : Génie Informatique", 13, True)
doc.add_paragraph()
add_centered("─" * 40, 12, color=RGBColor(0, 51, 102))
add_centered("Titre du projet :", 13)
add_centered("Conception et développement d'une application web\nde gestion immobilière pour Bakkah Immobilier", 16, True, RGBColor(0, 51, 102))
add_centered("─" * 40, 12, color=RGBColor(0, 51, 102))
doc.add_paragraph()
doc.add_paragraph()
add_centered("Réalisé par", 13, True)
add_centered("[Votre Nom et Prénom]", 14)
doc.add_paragraph()
add_centered("Encadré par", 13, True)
add_centered("Encadrant de l'école : [Nom du Professeur]", 12)
add_centered("Encadrant de l'entreprise : [Nom de l'encadrant]", 12)
doc.add_paragraph()
doc.add_paragraph()
add_centered("Année universitaire : 2025 - 2026", 14, True, RGBColor(0, 51, 102))
doc.add_paragraph().runs  # spacer
doc.add_page_break()

# ========== DÉDICACES ==========
doc.add_heading('Dédicaces', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
dedicace = """À MES CHERS PARENTS

Aucune dédicace ne saurait exprimer mon respect, mon amour éternel et ma considération pour les sacrifices que vous avez consentis pour mon instruction et mon bien-être. Je vous remercie pour tout le soutien et l'amour que vous me portez depuis mon enfance.

À MES FRÈRES ET SŒURS

En témoignage de mon affection fraternelle, de ma profonde tendresse et reconnaissance, je vous souhaite une vie pleine de bonheur et de succès.

À MES AMIS

En souvenir de notre sincère et profonde amitié et des moments agréables que nous avons passés ensemble. Veuillez trouver dans ce travail l'expression de mon respect le plus profond et mon affection la plus sincère.

À toute personne qui a contribué à mon éducation, à mes professeurs,
à toute personne m'ayant aidé un jour, à travers un conseil,
à tous ceux qui ont contribué de près ou de loin à la réussite de ce projet.

Que ce travail témoigne de mes sentiments les plus sincères."""

for line in dedicace.split('\n'):
    line = line.strip()
    if not line:
        doc.add_paragraph()
    elif line.startswith('À MES') or line.startswith('À toute') or line.startswith('Que ce'):
        add_para(line, bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        add_para(line, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ========== REMERCIEMENTS ==========
doc.add_heading('Remerciements', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
add_para("À Dieu, le tout puissant, nous rendons grâce pour nous avoir donné santé, patience, volonté et surtout raison.", italic=True)
add_para("Il m'est agréable d'exprimer ma reconnaissance auprès de toutes les personnes, dont l'intervention au cours de ce projet, de près ou de loin, a favorisé son aboutissement.")
add_para("En premier lieu, je remercie mes professeurs et l'ensemble des cadres administratifs de mon établissement pour la qualité de la formation qu'ils nous ont assurée, et tous les membres du jury qui m'ont fait l'honneur d'accepter de juger mon travail.")
add_para("Je tiens à exprimer ma profonde gratitude à mon encadrant de l'entreprise au sein de Bakkah Immobilier, pour sa disponibilité, pour le temps, et les conseils qu'il m'a prodigués, et qui m'ont été d'un fort appui.")
add_para("Je remercie mon encadrant pédagogique, pour ses conseils, son suivi, et pour tous les efforts qu'il a consentis tout au long de ma période de stage.")
add_para("Que tous ceux qui m'ont aidé et soutenu, de près ou de loin, trouvent ici l'expression de mes sentiments les plus distingués.", italic=True, bold=True)
doc.add_page_break()

# ========== AVANT-PROPOS ==========
doc.add_heading('Avant-Propos', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
add_para("Le présent travail intitulé « Conception et développement d'une application web de gestion immobilière pour Bakkah Immobilier » a été réalisé par l'élève ingénieur :")
add_para("[Votre Nom], étudiant(e) en Génie Informatique, dans le cadre du projet de fin d'études pour l'obtention du diplôme. Le projet a été réalisé en monôme.")
add_para("Ce stage a été encadré par [Nom du Professeur], enseignant à [Nom de l'établissement], et [Nom de l'encadrant entreprise], encadrant au sein de Bakkah Immobilier.")
add_para("Le stage a été effectué entre le [date début] et le [date fin].", italic=True)
doc.add_page_break()

# ========== RÉSUMÉ ==========
doc.add_heading('Résumé', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
add_para("Le présent rapport décrit le travail effectué durant notre stage de fin d'études au sein de la société Bakkah Immobilier, spécialisée dans la promotion immobilière au Maroc.")
add_para("L'objectif principal de ce projet est la conception et le développement d'une application web de gestion immobilière permettant le suivi en temps réel des projets de construction, la gestion des clients, l'automatisation des notifications, et l'intégration d'un assistant virtuel intelligent basé sur l'intelligence artificielle (IA).")
add_para("L'application a été développée en utilisant le framework Laravel 12 pour le backend (API RESTful), avec une architecture MVC, une authentification sécurisée via Laravel Sanctum, et une base de données MySQL. Le frontend a été conçu comme une Single Page Application (SPA) moderne.")
add_para("Mots-clés : Laravel, API RESTful, Gestion immobilière, Sanctum, MySQL, Intelligence Artificielle, GROQ API, Dashboard.", bold=True, italic=True)
doc.add_page_break()

# ========== TABLE DES MATIÈRES ==========
doc.add_heading('Table des Matières', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
toc_items = [
    ("Introduction Générale", ""),
    ("Chapitre 1 : Contexte Général du Projet", ""),
    ("   1.1 Présentation de l'organisme d'accueil", ""),
    ("   1.2 Contexte et problématique", ""),
    ("   1.3 Objectifs du projet", ""),
    ("   1.4 Méthodologie de travail", ""),
    ("Chapitre 2 : Analyse et Spécification des Besoins", ""),
    ("   2.1 Besoins fonctionnels", ""),
    ("   2.2 Besoins non fonctionnels", ""),
    ("   2.3 Diagramme de cas d'utilisation", ""),
    ("   2.4 Diagrammes de séquence", ""),
    ("Chapitre 3 : Conception", ""),
    ("   3.1 Architecture globale de l'application", ""),
    ("   3.2 Architecture MVC", ""),
    ("   3.3 Modèle de données (MCD / MLD)", ""),
    ("   3.4 Diagramme de classes", ""),
    ("Chapitre 4 : Réalisation", ""),
    ("   4.1 Environnement de développement", ""),
    ("   4.2 Technologies utilisées", ""),
    ("   4.3 Captures d'écran de l'application", ""),
    ("   4.4 Tests et validation", ""),
    ("Conclusion Générale et Perspectives", ""),
    ("Références et Webographie", ""),
    ("Annexes", ""),
]
for item, _ in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(item)
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    if not item.startswith("   "):
        r.bold = True

doc.add_page_break()

# ========== LISTE DES FIGURES ==========
doc.add_heading('Liste des Figures', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
figures = [
    "Figure 1 : Logo Bakkah Immobilier",
    "Figure 2 : Organigramme de Bakkah Immobilier",
    "Figure 3 : Diagramme de Gantt prévisionnel",
    "Figure 4 : Diagramme de Gantt réel",
    "Figure 5 : Diagramme de cas d'utilisation général",
    "Figure 6 : Diagramme de cas d'utilisation - Gestion des projets",
    "Figure 7 : Diagramme de séquence - Authentification",
    "Figure 8 : Diagramme de séquence - Gestion des projets",
    "Figure 9 : Diagramme de séquence - Chat IA",
    "Figure 10 : Architecture MVC de l'application",
    "Figure 11 : Modèle Conceptuel de Données (MCD)",
    "Figure 12 : Modèle Logique de Données (MLD)",
    "Figure 13 : Diagramme de classes",
    "Figure 14 : Architecture physique du système",
    "Figure 15 : Page d'authentification",
    "Figure 16 : Dashboard Client",
    "Figure 17 : Dashboard Administrateur",
    "Figure 18 : Gestion des projets",
    "Figure 19 : Détail d'un projet et étapes",
    "Figure 20 : Assistant virtuel IA",
    "Figure 21 : Formulaire de contact",
    "Figure 22 : Gestion des utilisateurs (Admin)",
]
for fig in figures:
    add_para(fig, size=11)
doc.add_page_break()

# ========== LISTE DES TABLEAUX ==========
doc.add_heading('Liste des Tableaux', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
tableaux = [
    "Tableau 1 : Planning prévisionnel du projet",
    "Tableau 2 : Besoins fonctionnels de l'application",
    "Tableau 3 : Besoins non fonctionnels",
    "Tableau 4 : Description des tables de la base de données",
    "Tableau 5 : Technologies et outils utilisés",
    "Tableau 6 : Points de terminaison de l'API RESTful",
]
for tab in tableaux:
    add_para(tab, size=11)
doc.add_page_break()

# ========== INTRODUCTION GÉNÉRALE ==========
doc.add_heading('Introduction Générale', level=1)
doc.add_paragraph()
add_para("Le secteur immobilier au Maroc connaît une croissance soutenue, portée par une demande croissante en logements et en espaces commerciaux. Les promoteurs immobiliers font face à des défis majeurs en matière de gestion de projets, de communication avec les clients et de suivi de l'avancement des travaux de construction.")
add_para("Dans ce contexte, Bakkah Immobilier, société marocaine spécialisée dans la promotion immobilière premium, a exprimé le besoin de disposer d'une plateforme web moderne permettant de centraliser la gestion de ses projets immobiliers et d'offrir à ses clients un accès transparent au suivi de leurs biens en cours de construction.")
add_para("C'est dans cette optique que s'inscrit notre projet de fin d'études, qui consiste à concevoir et développer une application web complète de gestion immobilière. Cette application intègre plusieurs fonctionnalités innovantes telles que :")
doc.add_paragraph("Le suivi en temps réel de l'avancement des projets avec des étapes détaillées", style='List Bullet')
doc.add_paragraph("Un tableau de bord administrateur avec statistiques et indicateurs clés", style='List Bullet')
doc.add_paragraph("Un système de notifications automatiques lors du changement de statut des projets", style='List Bullet')
doc.add_paragraph("Un assistant virtuel intelligent (chatbot IA) basé sur l'API GROQ/LLaMA", style='List Bullet')
doc.add_paragraph("L'export des données en formats CSV et PDF", style='List Bullet')
doc.add_paragraph("Un formulaire de contact avec envoi d'e-mails automatisé", style='List Bullet')
add_para("Le présent rapport s'articule autour de quatre chapitres principaux : le premier chapitre présente le contexte général du projet, le deuxième détaille l'analyse et les spécifications des besoins, le troisième est consacré à la conception, et le dernier chapitre couvre la phase de réalisation.")
doc.add_page_break()

# ========== CHAPITRE 1 ==========
doc.add_heading('Chapitre 1 : Contexte Général du Projet', level=1)

doc.add_heading('1.1 Présentation de l\'organisme d\'accueil', level=2)
add_para("Bakkah Immobilier est une société marocaine de promotion immobilière, présente depuis plus de 10 ans sur le marché. Elle se distingue par son engagement envers la qualité, l'innovation et la durabilité dans ses réalisations immobilières.")
add_para("Chiffres clés de Bakkah Immobilier :", bold=True)
doc.add_paragraph("Plus de 10 années d'expérience dans le secteur immobilier", style='List Bullet')
doc.add_paragraph("Plus de 500 logements livrés à ce jour", style='List Bullet')
doc.add_paragraph("15 projets en cours de réalisation", style='List Bullet')
doc.add_paragraph("98% de taux de satisfaction client", style='List Bullet')
add_para("Siège social : Appt.11 Rue Goumima Rés Le Louvre N°653 Etg.4, Casablanca, Maroc")
add_para("Valeurs fondamentales :", bold=True)
doc.add_paragraph("Qualité Sans Compromis : sélection des meilleurs matériaux et artisans", style='List Bullet')
doc.add_paragraph("Innovation : utilisation de technologies modernes pour un suivi transparent", style='List Bullet')
doc.add_paragraph("Durabilité : conception respectueuse de l'environnement", style='List Bullet')

doc.add_heading('1.2 Contexte et problématique', level=2)
add_para("Avant la mise en place de cette solution, la gestion des projets immobiliers chez Bakkah Immobilier s'effectuait de manière traditionnelle, ce qui engendrait plusieurs problèmes :")
doc.add_paragraph("Manque de visibilité des clients sur l'avancement de leurs projets", style='List Bullet')
doc.add_paragraph("Communication non structurée entre l'entreprise et ses clients", style='List Bullet')
doc.add_paragraph("Absence d'un système centralisé de suivi des étapes de construction", style='List Bullet')
doc.add_paragraph("Difficulté de générer des rapports et des statistiques en temps réel", style='List Bullet')
doc.add_paragraph("Pas de moyen automatisé de notification des clients lors de mises à jour", style='List Bullet')
add_para("La problématique peut se résumer ainsi : Comment concevoir et développer une solution web moderne, sécurisée et évolutive permettant à Bakkah Immobilier de gérer efficacement ses projets immobiliers tout en offrant aux clients une expérience de suivi transparente et interactive ?", bold=True, italic=True)

doc.add_heading('1.3 Objectifs du projet', level=2)
add_para("Les objectifs principaux de ce projet sont :")
doc.add_paragraph("Développer une API RESTful sécurisée avec Laravel 12 et Sanctum", style='List Bullet')
doc.add_paragraph("Implémenter un système d'authentification avec gestion des rôles (Admin/Client)", style='List Bullet')
doc.add_paragraph("Créer un module de gestion des projets avec suivi par étapes", style='List Bullet')
doc.add_paragraph("Mettre en place un système de notifications automatiques", style='List Bullet')
doc.add_paragraph("Intégrer un assistant virtuel IA utilisant l'API GROQ avec le modèle LLaMA", style='List Bullet')
doc.add_paragraph("Permettre l'export des données en CSV et PDF", style='List Bullet')
doc.add_paragraph("Créer un dashboard administrateur avec statistiques temps réel", style='List Bullet')
doc.add_paragraph("Implémenter un formulaire de contact avec envoi d'e-mails via SMTP", style='List Bullet')

doc.add_heading('1.4 Méthodologie de travail', level=2)
add_para("Pour mener à bien ce projet, nous avons adopté une approche itérative et incrémentale. Le développement s'est déroulé en plusieurs phases :")
doc.add_paragraph("Phase 1 : Analyse des besoins et rédaction du cahier des charges", style='List Bullet')
doc.add_paragraph("Phase 2 : Conception de l'architecture et modélisation UML", style='List Bullet')
doc.add_paragraph("Phase 3 : Développement du backend (API Laravel)", style='List Bullet')
doc.add_paragraph("Phase 4 : Développement du frontend (SPA)", style='List Bullet')
doc.add_paragraph("Phase 5 : Intégration, tests et déploiement", style='List Bullet')
doc.add_page_break()

# ========== CHAPITRE 2 ==========
doc.add_heading('Chapitre 2 : Analyse et Spécification des Besoins', level=1)

doc.add_heading('2.1 Besoins fonctionnels', level=2)
add_para("L'application doit répondre aux besoins fonctionnels suivants :")

# Tableau besoins fonctionnels
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Module'
hdr[1].text = 'Fonctionnalité'
hdr[2].text = 'Acteur'

bf = [
    ("Authentification", "Inscription, connexion, déconnexion sécurisée (Sanctum)", "Tous"),
    ("Gestion Profil", "Modifier nom, email, téléphone, mot de passe", "Client/Admin"),
    ("Gestion Projets", "CRUD projets, filtrage, recherche, upload documents", "Admin"),
    ("Suivi Étapes", "Ajouter/modifier/supprimer étapes, calcul auto du progrès", "Admin"),
    ("Dashboard Admin", "Statistiques : total projets, clients, CA, progrès moyen", "Admin"),
    ("Gestion Utilisateurs", "CRUD utilisateurs, attribution des rôles", "Admin"),
    ("Notifications", "Envoi automatique lors du changement de statut d'un projet", "Système"),
    ("Chat IA", "Assistant virtuel intelligent avec contexte des projets", "Client"),
    ("Contact", "Formulaire de contact avec envoi d'e-mail SMTP", "Client"),
    ("Export", "Export CSV et PDF des données des projets", "Admin/Client"),
    ("Paramètres", "Configuration dynamique de l'application (clé/valeur)", "Admin"),
]
for mod, fonc, act in bf:
    row = table.add_row().cells
    row[0].text = mod
    row[1].text = fonc
    row[2].text = act

doc.add_paragraph()

doc.add_heading('2.2 Besoins non fonctionnels', level=2)
doc.add_paragraph("Sécurité : Authentification par token (Sanctum), hachage des mots de passe (Bcrypt)", style='List Bullet')
doc.add_paragraph("Performance : Temps de réponse API < 500ms, mise en cache des requêtes", style='List Bullet')
doc.add_paragraph("Ergonomie : Interface intuitive, responsive design, expérience utilisateur premium", style='List Bullet')
doc.add_paragraph("Scalabilité : Architecture modulaire permettant l'ajout de nouvelles fonctionnalités", style='List Bullet')
doc.add_paragraph("Maintenabilité : Code structuré selon le pattern MVC, respect des conventions Laravel", style='List Bullet')
doc.add_paragraph("Fiabilité : Validation des données côté serveur, gestion des erreurs", style='List Bullet')

doc.add_heading('2.3 Diagramme de cas d\'utilisation', level=2)
add_para("[Insérer ici le diagramme de cas d'utilisation général]", italic=True, bold=True)
add_para("Le système distingue deux acteurs principaux :")
doc.add_paragraph("L'Administrateur : a accès à toutes les fonctionnalités de gestion", style='List Bullet')
doc.add_paragraph("Le Client : peut consulter ses projets, utiliser le chat IA et envoyer des messages", style='List Bullet')

doc.add_heading('2.4 Diagrammes de séquence', level=2)
add_para("2.4.1 Séquence d'authentification", bold=True)
add_para("[Insérer diagramme de séquence - Authentification]", italic=True)
add_para("Le processus d'authentification utilise Laravel Sanctum : l'utilisateur envoie ses identifiants (email/mot de passe) via une requête POST à /api/login. Le serveur vérifie les identifiants, crée un token d'accès personnel et le renvoie au client.")

add_para("2.4.2 Séquence de gestion des projets", bold=True)
add_para("[Insérer diagramme de séquence - Gestion projets]", italic=True)
add_para("L'administrateur authentifié peut créer, modifier et supprimer des projets via l'API RESTful. Chaque opération déclenche une validation côté serveur et peut générer une notification automatique vers le client concerné.")
doc.add_page_break()

# ========== CHAPITRE 3 ==========
doc.add_heading('Chapitre 3 : Conception', level=1)

doc.add_heading('3.1 Architecture globale de l\'application', level=2)
add_para("L'application suit une architecture client-serveur de type API RESTful :")
doc.add_paragraph("Backend : Laravel 12 (PHP 8.2) exposant une API REST", style='List Bullet')
doc.add_paragraph("Frontend : Application SPA communiquant avec l'API via Axios", style='List Bullet')
doc.add_paragraph("Base de données : MySQL (base 'stage')", style='List Bullet')
doc.add_paragraph("Authentification : Laravel Sanctum (tokens API)", style='List Bullet')
doc.add_paragraph("IA : API GROQ avec modèle LLaMA 3.1-8B", style='List Bullet')
doc.add_paragraph("E-mail : SMTP Gmail", style='List Bullet')

doc.add_heading('3.2 Architecture MVC', level=2)
add_para("Laravel implémente le pattern MVC (Modèle-Vue-Contrôleur) :")
add_para("Modèles (5) : User, Project, ProjectStep, Notification, Setting — définissent la structure des données et les relations (hasMany, belongsTo).", bold=False)
add_para("Contrôleurs (7) : AuthController, AdminController, ProjectController, ProjectStepController, ChatController, ContactController, SettingController — contiennent la logique métier.", bold=False)
add_para("Vues : Templates Blade pour les e-mails et les exports PDF.", bold=False)

doc.add_heading('3.3 Modèle de données', level=2)
add_para("La base de données comprend les tables suivantes :", bold=True)

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Table'
hdr2[1].text = 'Colonnes principales'
hdr2[2].text = 'Relations'

tables_data = [
    ("users", "id, name, email, password, role, avatar, phone", "hasMany(Project, Notification)"),
    ("projects", "id, user_id, name, location, progress, status, image_path, document_path, total_value, delivery_date", "belongsTo(User), hasMany(ProjectStep)"),
    ("project_steps", "id, project_id, label, status (enum), order_num", "belongsTo(Project)"),
    ("notifications", "id, user_id, content, is_read", "belongsTo(User)"),
    ("settings", "id, key, value", "Aucune"),
]
for tbl, cols, rels in tables_data:
    row = table2.add_row().cells
    row[0].text = tbl
    row[1].text = cols
    row[2].text = rels

doc.add_paragraph()

doc.add_heading('3.4 Diagramme de classes', level=2)
add_para("[Insérer le diagramme de classes UML]", italic=True, bold=True)
add_para("Les principales relations entre les classes sont :")
doc.add_paragraph("User (1) ──── (*) Project : Un utilisateur possède plusieurs projets", style='List Bullet')
doc.add_paragraph("Project (1) ──── (*) ProjectStep : Un projet contient plusieurs étapes", style='List Bullet')
doc.add_paragraph("User (1) ──── (*) Notification : Un utilisateur reçoit plusieurs notifications", style='List Bullet')
doc.add_page_break()

# ========== CHAPITRE 4 ==========
doc.add_heading('Chapitre 4 : Réalisation', level=1)

doc.add_heading('4.1 Environnement de développement', level=2)
table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Light Grid Accent 1'
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Outil / Technologie'
hdr3[1].text = 'Version / Description'

env_data = [
    ("Système d'exploitation", "Windows"),
    ("Serveur local", "WAMP64 (Apache, MySQL, PHP)"),
    ("Éditeur de code", "Visual Studio Code"),
    ("PHP", "8.2+"),
    ("Laravel", "12.0"),
    ("Composer", "Gestionnaire de dépendances PHP"),
    ("Node.js / NPM", "Runtime JavaScript + gestionnaire de paquets"),
    ("Vite", "7.0+ (bundler frontend)"),
    ("MySQL", "Base de données relationnelle"),
    ("Git", "Contrôle de version"),
    ("Postman", "Test des API REST"),
]
for outil, desc in env_data:
    row = table3.add_row().cells
    row[0].text = outil
    row[1].text = desc

doc.add_paragraph()

doc.add_heading('4.2 Technologies utilisées', level=2)
add_para("Backend :", bold=True)
doc.add_paragraph("Laravel 12 : Framework PHP pour le développement de l'API RESTful", style='List Bullet')
doc.add_paragraph("Laravel Sanctum 4.0 : Authentification par tokens API (SPA-friendly)", style='List Bullet')
doc.add_paragraph("Eloquent ORM : Gestion des modèles et des relations de base de données", style='List Bullet')
doc.add_paragraph("DomPDF (barryvdh/laravel-dompdf 3.1) : Génération et export de rapports PDF", style='List Bullet')
doc.add_paragraph("API GROQ : Intégration IA avec modèle LLaMA 3.1-8B pour le chatbot", style='List Bullet')

add_para("Frontend :", bold=True)
doc.add_paragraph("HTML5 / CSS3 / JavaScript : Technologies web standards", style='List Bullet')
doc.add_paragraph("TailwindCSS 4.0 : Framework CSS utilitaire pour un design moderne", style='List Bullet')
doc.add_paragraph("Axios : Client HTTP pour la communication avec l'API", style='List Bullet')
doc.add_paragraph("Vite 7 : Outil de build et serveur de développement rapide", style='List Bullet')

doc.add_heading('4.3 Points de terminaison de l\'API (Routes)', level=2)
add_para("L'API RESTful expose les endpoints suivants :")

table4 = doc.add_table(rows=1, cols=4)
table4.style = 'Light Grid Accent 1'
hdr4 = table4.rows[0].cells
hdr4[0].text = 'Méthode'
hdr4[1].text = 'Endpoint'
hdr4[2].text = 'Description'
hdr4[3].text = 'Auth'

routes = [
    ("POST", "/api/register", "Inscription d'un nouveau client", "Non"),
    ("POST", "/api/login", "Connexion et obtention du token", "Non"),
    ("POST", "/api/logout", "Déconnexion (suppression du token)", "Oui"),
    ("GET", "/api/user", "Récupérer le profil de l'utilisateur connecté", "Oui"),
    ("PUT", "/api/user/profile", "Modifier le profil", "Oui"),
    ("PUT", "/api/user/password", "Changer le mot de passe", "Oui"),
    ("GET", "/api/admin/stats", "Statistiques du dashboard admin", "Oui"),
    ("GET/POST/PUT/DELETE", "/api/admin/users", "CRUD des utilisateurs", "Oui"),
    ("GET", "/api/projects", "Liste des projets (filtrable)", "Oui"),
    ("POST", "/api/projects", "Créer un nouveau projet", "Oui"),
    ("PUT", "/api/projects/{id}", "Modifier un projet", "Oui"),
    ("DELETE", "/api/projects/{id}", "Supprimer un projet", "Oui"),
    ("GET", "/api/projects/export", "Export CSV", "Oui"),
    ("GET", "/api/projects/export-pdf", "Export PDF", "Oui"),
    ("POST/PUT/DELETE", "/api/steps", "CRUD des étapes de projet", "Oui"),
    ("POST", "/api/chat", "Envoyer un message au chatbot IA", "Non"),
    ("POST", "/api/contact", "Envoyer un e-mail de contact", "Oui"),
    ("GET/POST", "/api/settings", "Gestion des paramètres", "Oui"),
]
for meth, ep, desc, auth in routes:
    row = table4.add_row().cells
    row[0].text = meth
    row[1].text = ep
    row[2].text = desc
    row[3].text = auth

doc.add_paragraph()

doc.add_heading('4.4 Captures d\'écran de l\'application', level=2)
add_para("[Insérer les captures d'écran des différentes pages de l'application]", italic=True, bold=True)
add_para("Page de connexion / inscription", bold=True)
add_para("La page d'authentification permet aux utilisateurs de se connecter avec leur email et mot de passe, ou de créer un nouveau compte client.")
add_para("Dashboard Client", bold=True)
add_para("Le tableau de bord client affiche la liste des projets attribués avec leur pourcentage d'avancement, statut actuel, et détails des étapes de construction.")
add_para("Dashboard Administrateur", bold=True)
add_para("Le dashboard admin présente les statistiques clés : nombre total de projets, de clients, chiffre d'affaires, progrès moyen, ainsi que les activités récentes.")
add_para("Module Chat IA", bold=True)
add_para("L'assistant virtuel intelligent, alimenté par l'API GROQ et le modèle LLaMA 3.1, permet aux clients de poser des questions sur leurs projets et sur Bakkah Immobilier en temps réel.")
add_para("Gestion des utilisateurs", bold=True)
add_para("L'interface d'administration permet de créer, modifier et supprimer des comptes utilisateurs avec attribution des rôles (admin/client).")

doc.add_heading('4.5 Tests et validation', level=2)
add_para("Les tests ont été effectués à plusieurs niveaux :")
doc.add_paragraph("Tests unitaires : Validation des modèles et des contrôleurs", style='List Bullet')
doc.add_paragraph("Tests d'intégration : Vérification des endpoints API avec Postman", style='List Bullet')
doc.add_paragraph("Tests fonctionnels : Validation des scénarios utilisateur de bout en bout", style='List Bullet')
doc.add_paragraph("Tests de sécurité : Vérification de l'authentification Sanctum et des autorisations", style='List Bullet')
doc.add_page_break()

# ========== CONCLUSION ==========
doc.add_heading('Conclusion Générale et Perspectives', level=1)
doc.add_paragraph()
add_para("Au terme de ce projet de fin d'études, nous avons réussi à concevoir et développer une application web complète de gestion immobilière pour la société Bakkah Immobilier. Cette application répond aux besoins exprimés en offrant une solution moderne, sécurisée et performante.")
add_para("Les principales réalisations incluent :")
doc.add_paragraph("Une API RESTful robuste avec Laravel 12 et une authentification sécurisée via Sanctum", style='List Bullet')
doc.add_paragraph("Un système de gestion de projets avec suivi par étapes et calcul automatique du progrès", style='List Bullet')
doc.add_paragraph("Un dashboard administrateur avec statistiques en temps réel", style='List Bullet')
doc.add_paragraph("Un assistant virtuel IA innovant basé sur l'API GROQ/LLaMA", style='List Bullet')
doc.add_paragraph("Un système de notifications automatiques et d'envoi d'e-mails", style='List Bullet')
doc.add_paragraph("Des fonctionnalités d'export en CSV et PDF", style='List Bullet')

add_para("Perspectives d'amélioration :", bold=True)
doc.add_paragraph("Ajout d'un système de messagerie en temps réel (WebSocket)", style='List Bullet')
doc.add_paragraph("Intégration de la géolocalisation des projets sur une carte interactive", style='List Bullet')
doc.add_paragraph("Développement d'une application mobile (React Native ou Flutter)", style='List Bullet')
doc.add_paragraph("Mise en place d'un système de paiement en ligne", style='List Bullet')
doc.add_paragraph("Amélioration du chatbot IA avec un historique de conversations", style='List Bullet')
doc.add_paragraph("Déploiement en production sur un serveur cloud (AWS, DigitalOcean)", style='List Bullet')

add_para("Ce stage nous a permis de consolider nos connaissances en développement web full-stack, de maîtriser le framework Laravel et ses écosystèmes, et de découvrir les possibilités offertes par l'intégration de l'intelligence artificielle dans les applications web modernes.", italic=True)
doc.add_page_break()

# ========== RÉFÉRENCES ==========
doc.add_heading('Références et Webographie', level=1)
doc.add_paragraph()
refs = [
    "Laravel Documentation officielle : https://laravel.com/docs",
    "Laravel Sanctum : https://laravel.com/docs/sanctum",
    "GROQ API Documentation : https://console.groq.com/docs",
    "MySQL Documentation : https://dev.mysql.com/doc/",
    "TailwindCSS Documentation : https://tailwindcss.com/docs",
    "Vite Documentation : https://vitejs.dev/guide/",
    "DomPDF pour Laravel : https://github.com/barryvdh/laravel-dompdf",
    "Axios Documentation : https://axios-http.com/docs/intro",
    "PHP.net Documentation : https://www.php.net/docs.php",
]
for i, ref in enumerate(refs, 1):
    add_para(f"[{i}] {ref}", size=11)
doc.add_page_break()

# ========== ANNEXES ==========
doc.add_heading('Annexes', level=1)
doc.add_paragraph()
add_para("Annexe A : Code source des principaux contrôleurs", bold=True)
add_para("Le code source complet est disponible dans le répertoire du projet. Les contrôleurs principaux sont :")
doc.add_paragraph("AuthController.php : Gestion de l'authentification (register, login, logout, updateProfile, updatePassword)", style='List Bullet')
doc.add_paragraph("ProjectController.php : CRUD des projets, export CSV/PDF", style='List Bullet')
doc.add_paragraph("ProjectStepController.php : Gestion des étapes avec recalcul automatique du progrès", style='List Bullet')
doc.add_paragraph("AdminController.php : Dashboard admin, CRUD utilisateurs", style='List Bullet')
doc.add_paragraph("ChatController.php : Intégration IA avec GROQ API et contexte dynamique", style='List Bullet')
doc.add_paragraph("ContactController.php : Envoi de mails via SMTP", style='List Bullet')
doc.add_paragraph("SettingController.php : Gestion des paramètres clé/valeur", style='List Bullet')

doc.add_paragraph()
add_para("Annexe B : Structure de la base de données (Migrations)", bold=True)
add_para("Les fichiers de migration définissent le schéma de la base de données 'stage' sous MySQL, avec les tables : users, projects, project_steps, notifications, settings, personal_access_tokens, cache, jobs.")

# ---- Save ----
output_path = r'c:\wamp64\www\pfe-stage-laravel1\Rapport_de_Stage_Bakkah_Immobilier.docx'
doc.save(output_path)
print(f"Rapport généré avec succès : {output_path}")
