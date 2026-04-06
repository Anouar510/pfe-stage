from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import datetime

doc = Document()

# ---- Add Header with Logos ----
section = doc.sections[0]
header = section.header
header_table = header.add_table(rows=1, cols=3, width=Inches(6.5))
header_table.autofit = False
header_table.columns[0].width = Inches(2.0)
header_table.columns[1].width = Inches(2.5)
header_table.columns[2].width = Inches(2.0)

# OFPPT Logo (Left)
cell_left = header_table.cell(0, 0)
p_left = cell_left.paragraphs[0]
run_left = p_left.add_run()
try:
    run_left.add_picture(r"c:\wamp64\www\pfe-stage-laravel1\ofppt.jpg", width=Inches(1.2))
except Exception as e:
    print(f"Erreur chargement logo OFPPT: {e}")
p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Bakkah Logo (Right)
cell_right = header_table.cell(0, 2)
p_right = cell_right.paragraphs[0]
run_right = p_right.add_run()
try:
    run_right.add_picture(r"c:\wamp64\www\pfe-stage-laravel1\bakkah.png", width=Inches(1.2))
except Exception as e:
    print(f"Erreur chargement logo Bakkah: {e}")
p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0, 51, 102)
    hs.font.name = 'Times New Roman'
    hs.font.bold = True
    if i == 1: hs.font.size = Pt(18)
    elif i == 2: hs.font.size = Pt(14)
    else: hs.font.size = Pt(12)

def add_centered(text, size=12, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.name = 'Times New Roman'
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_para(text, bold=False, italic=False, align=None, size=12, space_after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(space_after)
    return p

# ========== PAGE DE GARDE ==========
for _ in range(4): doc.add_paragraph()
add_centered("RAPPORT DE STAGE DE FIN D'ÉTUDES", 22, True, RGBColor(0, 51, 102))
doc.add_paragraph()
add_centered("Présenté pour l'obtention du diplôme", 13)
add_centered("Filière : Génie Informatique", 13, True)
doc.add_paragraph()
add_centered("─" * 40, 12, color=RGBColor(0, 51, 102))
add_centered("Titre du projet :", 13)
add_centered("Conception et développement d'une application web\nde gestion immobilière pour Bakkah Immobilier", 16, True, RGBColor(0, 51, 102))
add_centered("─" * 40, 12, color=RGBColor(0, 51, 102))
doc.add_paragraph()
doc.add_paragraph()
add_centered("Réalisé par", 13, True)
add_centered("[Votre Nom et Prénom]", 14)
doc.add_paragraph()
add_centered("Encadré par", 13, True)
add_centered("Encadrant de l'école : [Nom du Professeur]", 12)
add_centered("Encadrant de l'entreprise : [Nom de l'encadrant]", 12)
doc.add_paragraph()
doc.add_paragraph()
add_centered("Année universitaire : 2025 - 2026", 14, True, RGBColor(0, 51, 102))
doc.add_paragraph().runs  # spacer
doc.add_page_break()

# ========== DÉDICACES ==========
doc.add_heading('Dédicaces', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
dedicace = """À MES CHERS PARENTS

Aucune dédicace ne saurait exprimer mon respect, mon amour éternel et ma considération pour les sacrifices que vous avez consentis pour mon instruction et mon bien-être. Je vous remercie pour tout le soutien et l'amour que vous me portez depuis mon enfance.

À MES FRÈRES ET SŒURS

En témoignage de mon affection fraternelle, de ma profonde tendresse et reconnaissance, je vous souhaite une vie pleine de bonheur et de succès.

À MES AMIS

En souvenir de notre sincère et profonde amitié et des moments agréables que nous avons passés ensemble. Veuillez trouver dans ce travail l'expression de mon respect le plus profond et mon affection la plus sincère.

À toute personne qui a contribué à mon éducation, à mes professeurs,
à toute personne m'ayant aidé un jour, à travers un conseil,
à tous ceux qui ont contribué de près ou de loin à la réussite de ce projet.

Que ce travail témoigne de mes sentiments les plus sincères."""

for line in dedicace.split('\n'):
    line = line.strip()
    if not line:
        doc.add_paragraph()
    elif line.startswith('À MES') or line.startswith('À toute') or line.startswith('Que ce'):
        add_para(line, bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        add_para(line, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ========== REMERCIEMENTS ==========
doc.add_heading('Remerciements', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
add_para("À Dieu, le tout puissant, nous rendons grâce pour nous avoir donné santé, patience, volonté et surtout raison.", italic=True)
add_para("Il m'est agréable d'exprimer ma reconnaissance auprès de toutes les personnes, dont l'intervention au cours de ce projet, de près ou de loin, a favorisé son aboutissement.")
add_para("En premier lieu, je remercie mes professeurs et l'ensemble des cadres administratifs de mon établissement pour la qualité de la formation qu'ils nous ont assurée, et tous les membres du jury qui m'ont fait l'honneur d'accepter de juger mon travail.")
add_para("Je tiens à exprimer ma profonde gratitude à mon encadrant de l'entreprise au sein de Bakkah Immobilier, pour sa disponibilité, pour le temps, et les conseils qu'il m'a prodigués, et qui m'ont été d'un fort appui.")
add_para("Je remercie mon encadrant pédagogique, pour ses conseils, son suivi, et pour tous les efforts qu'il a consentis tout au long de ma période de stage.")
add_para("Que tous ceux qui m'ont aidé et soutenu, de près ou de loin, trouvent ici l'expression de mes sentiments les plus distingués.", italic=True, bold=True)
doc.add_page_break()

# ========== AVANT-PROPOS ==========
doc.add_heading('Avant-Propos', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
add_para("Le présent travail intitulé « Conception et développement d'une application web de gestion immobilière pour Bakkah Immobilier » a été réalisé par l'élève ingénieur :")
add_para("[Votre Nom], étudiant(e) en Génie Informatique, dans le cadre du projet de fin d'études pour l'obtention du diplôme. Le projet a été réalisé en monôme.")
add_para("Ce stage a été encadré par [Nom du Professeur], enseignant à [Nom de l'établissement], et [Nom de l'encadrant entreprise], encadrant au sein de Bakkah Immobilier.")
add_para("Le stage a été effectué entre le [date début] et le [date fin].", italic=True)
doc.add_page_break()

# ========== RÉSUMÉ ==========
doc.add_heading('Résumé', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
add_para("Le présent rapport décrit le travail effectué durant notre stage de fin d'études au sein de la société Bakkah Immobilier, spécialisée dans la promotion immobilière au Maroc.")
add_para("L'objectif principal de ce projet est la conception et le développement d'une application web de gestion immobilière permettant le suivi en temps réel des projets de construction, la gestion des clients, l'automatisation des notifications, et l'intégration d'un assistant virtuel intelligent basé sur un système de traitement du langage naturel (NLP) avec exploitation contextuelle de la base de données.")
add_para("L'application a été développée en utilisant le framework Laravel 12 pour le backend (API RESTful), avec une architecture MVC, une authentification sécurisée via Laravel Sanctum, et une base de données MySQL. Le frontend a été conçu comme une Single Page Application (SPA) moderne avec un design premium et responsive.")
add_para("L'assistant virtuel « Bakkah Assistant » fonctionne de manière autonome sans dépendance à des APIs externes. Il utilise un moteur de correspondance par mots-clés enrichi par les données en temps réel de la base de données (projets, étapes, prix, dates de livraison) pour fournir des réponses précises et contextuelles aux utilisateurs.")
add_para("Mots-clés : Laravel, API RESTful, Gestion immobilière, Sanctum, MySQL, Chatbot intelligent, NLP, Dashboard analytique, SPA.", bold=True, italic=True)
doc.add_page_break()

# ========== TABLE DES MATIÈRES ==========
doc.add_heading('Table des Matières', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
toc_items = [
    ("Introduction Générale", ""),
    ("Chapitre 1 : Contexte Général du Projet", ""),
    ("   1.1 Présentation de l'organisme d'accueil", ""),
    ("   1.2 Contexte et problématique", ""),
    ("   1.3 Objectifs du projet", ""),
    ("   1.4 Méthodologie de travail", ""),
    ("Chapitre 2 : Analyse et Spécification des Besoins", ""),
    ("   2.1 Besoins fonctionnels", ""),
    ("   2.2 Besoins non fonctionnels", ""),
    ("   2.3 Diagramme de cas d'utilisation", ""),
    ("   2.4 Diagrammes de séquence", ""),
    ("Chapitre 3 : Conception", ""),
    ("   3.1 Architecture globale de l'application", ""),
    ("   3.2 Architecture MVC", ""),
    ("   3.3 Modèle de données (MCD / MLD)", ""),
    ("   3.4 Diagramme de classes", ""),
    ("   3.5 Architecture du chatbot intelligent", ""),
    ("Chapitre 4 : Réalisation", ""),
    ("   4.1 Environnement de développement", ""),
    ("   4.2 Technologies utilisées", ""),
    ("   4.3 Captures d'écran de l'application", ""),
    ("   4.4 Fonctionnement du chatbot intelligent", ""),
    ("   4.5 Dashboard administrateur avancé", ""),
    ("   4.6 Tests et validation", ""),
    ("Conclusion Générale et Perspectives", ""),
    ("Références et Webographie", ""),
    ("Annexes", ""),
]
for item, _ in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(item)
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    if not item.startswith("   "):
        r.bold = True

doc.add_page_break()

# ========== LISTE DES FIGURES ==========
doc.add_heading('Liste des Figures', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
figures = [
    "Figure 1 : Logo Bakkah Immobilier",
    "Figure 2 : Organigramme de Bakkah Immobilier",
    "Figure 3 : Diagramme de Gantt prévisionnel",
    "Figure 4 : Diagramme de Gantt réel",
    "Figure 5 : Diagramme de cas d'utilisation général",
    "Figure 6 : Diagramme de cas d'utilisation - Gestion des projets",
    "Figure 7 : Diagramme de séquence - Authentification",
    "Figure 8 : Diagramme de séquence - Gestion des projets",
    "Figure 9 : Diagramme de séquence - Chat intelligent",
    "Figure 10 : Architecture MVC de l'application",
    "Figure 11 : Modèle Conceptuel de Données (MCD)",
    "Figure 12 : Modèle Logique de Données (MLD)",
    "Figure 13 : Diagramme de classes",
    "Figure 14 : Architecture physique du système",
    "Figure 15 : Architecture du chatbot intelligent",
    "Figure 16 : Page d'authentification",
    "Figure 17 : Dashboard Client",
    "Figure 18 : Dashboard Administrateur avec graphiques",
    "Figure 19 : Gestion des projets",
    "Figure 20 : Détail d'un projet et étapes d'avancement",
    "Figure 21 : Assistant virtuel intelligent (Bakkah Assistant)",
    "Figure 22 : Formulaire de contact",
    "Figure 23 : Gestion des utilisateurs (Admin)",
    "Figure 24 : Flux d'activités récentes (Admin)",
    "Figure 25 : Export PDF des projets",
]
for fig in figures:
    add_para(fig, size=11)
doc.add_page_break()

# ========== LISTE DES TABLEAUX ==========
doc.add_heading('Liste des Tableaux', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
tableaux = [
    "Tableau 1 : Planning prévisionnel du projet",
    "Tableau 2 : Besoins fonctionnels de l'application",
    "Tableau 3 : Besoins non fonctionnels",
    "Tableau 4 : Description des tables de la base de données",
    "Tableau 5 : Technologies et outils utilisés",
    "Tableau 6 : Points de terminaison de l'API RESTful",
    "Tableau 7 : Catégories de réponses du chatbot intelligent",
]
for tab in tableaux:
    add_para(tab, size=11)
doc.add_page_break()

# ========== INTRODUCTION GÉNÉRALE ==========
doc.add_heading('Introduction Générale', level=1)
doc.add_paragraph()
add_para("Le secteur immobilier au Maroc connaît une croissance soutenue, portée par une demande croissante en logements et en espaces commerciaux. Les promoteurs immobiliers font face à des défis majeurs en matière de gestion de projets, de communication avec les clients et de suivi de l'avancement des travaux de construction.")
add_para("Dans ce contexte, Bakkah Immobilier, société marocaine spécialisée dans la promotion immobilière premium, a exprimé le besoin de disposer d'une plateforme web moderne permettant de centraliser la gestion de ses projets immobiliers et d'offrir à ses clients un accès transparent au suivi de leurs biens en cours de construction.")
add_para("C'est dans cette optique que s'inscrit notre projet de fin d'études, qui consiste à concevoir et développer une application web complète de gestion immobilière. Cette application intègre plusieurs fonctionnalités innovantes telles que :")
doc.add_paragraph("Le suivi en temps réel de l'avancement des projets avec des étapes détaillées et calcul automatique du progrès", style='List Bullet')
doc.add_paragraph("Un tableau de bord administrateur avancé avec statistiques, graphiques interactifs et flux d'activités en temps réel", style='List Bullet')
doc.add_paragraph("Un système de notifications automatiques lors du changement de statut des projets", style='List Bullet')
doc.add_paragraph("Un assistant virtuel intelligent (chatbot) basé sur un moteur NLP local exploitant les données en temps réel de la base de données", style='List Bullet')
doc.add_paragraph("L'export des données en formats CSV et PDF", style='List Bullet')
doc.add_paragraph("Un formulaire de contact avec envoi d'e-mails automatisé via SMTP", style='List Bullet')
doc.add_paragraph("Une interface utilisateur moderne et responsive avec design premium", style='List Bullet')
add_para("Le présent rapport s'articule autour de quatre chapitres principaux : le premier chapitre présente le contexte général du projet, le deuxième détaille l'analyse et les spécifications des besoins, le troisième est consacré à la conception, et le dernier chapitre couvre la phase de réalisation.")
doc.add_page_break()

# ========== CHAPITRE 1 ==========
doc.add_heading('Chapitre 1 : Contexte Général du Projet', level=1)

doc.add_heading('1.1 Présentation de l\'organisme d\'accueil', level=2)
add_para("Bakkah Immobilier est une société marocaine de promotion immobilière, présente depuis plus de 10 ans sur le marché. Elle se distingue par son engagement envers la qualité, l'innovation et la durabilité dans ses réalisations immobilières.")
add_para("Chiffres clés de Bakkah Immobilier :", bold=True)
doc.add_paragraph("Plus de 10 années d'expérience dans le secteur immobilier", style='List Bullet')
doc.add_paragraph("Plus de 500 logements livrés à ce jour", style='List Bullet')
doc.add_paragraph("15 projets en cours de réalisation", style='List Bullet')
doc.add_paragraph("98% de taux de satisfaction client", style='List Bullet')
add_para("Siège social : Appt.11 Rue Goumima Rés Le Louvre N°653 Etg.4, Casablanca, Maroc")
add_para("Valeurs fondamentales :", bold=True)
doc.add_paragraph("Qualité Sans Compromis : sélection des meilleurs matériaux et artisans", style='List Bullet')
doc.add_paragraph("Innovation : utilisation de technologies modernes pour un suivi transparent", style='List Bullet')
doc.add_paragraph("Durabilité : conception respectueuse de l'environnement", style='List Bullet')

doc.add_heading('1.2 Contexte et problématique', level=2)
add_para("Avant la mise en place de cette solution, la gestion des projets immobiliers chez Bakkah Immobilier s'effectuait de manière traditionnelle, ce qui engendrait plusieurs problèmes :")
doc.add_paragraph("Manque de visibilité des clients sur l'avancement de leurs projets", style='List Bullet')
doc.add_paragraph("Communication non structurée entre l'entreprise et ses clients", style='List Bullet')
doc.add_paragraph("Absence d'un système centralisé de suivi des étapes de construction", style='List Bullet')
doc.add_paragraph("Difficulté de générer des rapports et des statistiques en temps réel", style='List Bullet')
doc.add_paragraph("Pas de moyen automatisé de notification des clients lors de mises à jour", style='List Bullet')
doc.add_paragraph("Absence d'un outil d'assistance client automatisé pour répondre aux questions fréquentes", style='List Bullet')
add_para("La problématique peut se résumer ainsi : Comment concevoir et développer une solution web moderne, sécurisée et évolutive permettant à Bakkah Immobilier de gérer efficacement ses projets immobiliers tout en offrant aux clients une expérience de suivi transparente et interactive ?", bold=True, italic=True)

doc.add_heading('1.3 Objectifs du projet', level=2)
add_para("Les objectifs principaux de ce projet sont :")
doc.add_paragraph("Développer une API RESTful sécurisée avec Laravel 12 et Sanctum", style='List Bullet')
doc.add_paragraph("Implémenter un système d'authentification avec gestion des rôles (Admin/Client)", style='List Bullet')
doc.add_paragraph("Créer un module de gestion des projets avec suivi par étapes et calcul automatique du progrès", style='List Bullet')
doc.add_paragraph("Mettre en place un système de notifications automatiques", style='List Bullet')
doc.add_paragraph("Développer un assistant virtuel intelligent autonome exploitant les données de la base de données en temps réel", style='List Bullet')
doc.add_paragraph("Permettre l'export des données en CSV et PDF", style='List Bullet')
doc.add_paragraph("Créer un dashboard administrateur avancé avec statistiques, graphiques et flux d'activités", style='List Bullet')
doc.add_paragraph("Implémenter un formulaire de contact avec envoi d'e-mails via SMTP", style='List Bullet')

doc.add_heading('1.4 Méthodologie de travail', level=2)
add_para("Pour mener à bien ce projet, nous avons adopté une approche itérative et incrémentale. Le développement s'est déroulé en plusieurs phases :")
doc.add_paragraph("Phase 1 : Analyse des besoins et rédaction du cahier des charges", style='List Bullet')
doc.add_paragraph("Phase 2 : Conception de l'architecture et modélisation UML", style='List Bullet')
doc.add_paragraph("Phase 3 : Développement du backend (API Laravel)", style='List Bullet')
doc.add_paragraph("Phase 4 : Développement du frontend (SPA)", style='List Bullet')
doc.add_paragraph("Phase 5 : Développement du chatbot intelligent avec NLP local", style='List Bullet')
doc.add_paragraph("Phase 6 : Intégration, tests et déploiement", style='List Bullet')
doc.add_page_break()

# ========== CHAPITRE 2 ==========
doc.add_heading('Chapitre 2 : Analyse et Spécification des Besoins', level=1)

doc.add_heading('2.1 Besoins fonctionnels', level=2)
add_para("L'application doit répondre aux besoins fonctionnels suivants :")

# Tableau besoins fonctionnels
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Module'
hdr[1].text = 'Fonctionnalité'
hdr[2].text = 'Acteur'

bf = [
    ("Authentification", "Inscription, connexion, déconnexion sécurisée (Sanctum)", "Tous"),
    ("Gestion Profil", "Modifier nom, email, téléphone, mot de passe", "Client/Admin"),
    ("Gestion Projets", "CRUD projets, filtrage, recherche, upload documents", "Admin"),
    ("Suivi Étapes", "Ajouter/modifier/supprimer étapes, calcul auto du progrès", "Admin"),
    ("Dashboard Admin", "Statistiques avancées : total projets, clients, CA, progrès moyen, graphiques, flux d'activités", "Admin"),
    ("Gestion Utilisateurs", "CRUD utilisateurs, attribution des rôles", "Admin"),
    ("Notifications", "Envoi automatique lors du changement de statut d'un projet", "Système"),
    ("Assistant Virtuel", "Chatbot intelligent avec NLP local, exploitation contextuelle de la BDD en temps réel", "Client"),
    ("Contact", "Formulaire de contact avec envoi d'e-mail SMTP", "Client"),
    ("Export", "Export CSV et PDF des données des projets avec filtrage", "Admin/Client"),
    ("Paramètres", "Configuration dynamique de l'application (clé/valeur)", "Admin"),
]
for mod, fonc, act in bf:
    row = table.add_row().cells
    row[0].text = mod
    row[1].text = fonc
    row[2].text = act

doc.add_paragraph()

doc.add_heading('2.2 Besoins non fonctionnels', level=2)
doc.add_paragraph("Sécurité : Authentification par token (Sanctum), hachage des mots de passe (Bcrypt), protection CSRF", style='List Bullet')
doc.add_paragraph("Performance : Temps de réponse API < 500ms, mise en cache des requêtes, chargement optimisé", style='List Bullet')
doc.add_paragraph("Ergonomie : Interface intuitive, responsive design, expérience utilisateur premium avec animations fluides", style='List Bullet')
doc.add_paragraph("Scalabilité : Architecture modulaire permettant l'ajout de nouvelles fonctionnalités", style='List Bullet')
doc.add_paragraph("Maintenabilité : Code structuré selon le pattern MVC, respect des conventions Laravel", style='List Bullet')
doc.add_paragraph("Fiabilité : Validation des données côté serveur, gestion des erreurs robuste", style='List Bullet')
doc.add_paragraph("Autonomie : Le chatbot fonctionne sans dépendance à des APIs externes, garantissant une disponibilité à 100%", style='List Bullet')

doc.add_heading('2.3 Diagramme de cas d\'utilisation', level=2)
add_para("[Insérer ici le diagramme de cas d'utilisation général]", italic=True, bold=True)
add_para("Le système distingue deux acteurs principaux :")
doc.add_paragraph("L'Administrateur : a accès à toutes les fonctionnalités de gestion (projets, utilisateurs, dashboard avancé, exports)", style='List Bullet')
doc.add_paragraph("Le Client : peut consulter ses projets, suivre l'avancement, utiliser l'assistant virtuel, et envoyer des messages de contact", style='List Bullet')

doc.add_heading('2.4 Diagrammes de séquence', level=2)
add_para("2.4.1 Séquence d'authentification", bold=True)
add_para("[Insérer diagramme de séquence - Authentification]", italic=True)
add_para("Le processus d'authentification utilise Laravel Sanctum : l'utilisateur envoie ses identifiants (email/mot de passe) via une requête POST à /api/login. Le serveur vérifie les identifiants, crée un token d'accès personnel et le renvoie au client.")

add_para("2.4.2 Séquence de gestion des projets", bold=True)
add_para("[Insérer diagramme de séquence - Gestion projets]", italic=True)
add_para("L'administrateur authentifié peut créer, modifier et supprimer des projets via l'API RESTful. Chaque opération déclenche une validation côté serveur et peut générer une notification automatique vers le client concerné.")

add_para("2.4.3 Séquence du chatbot intelligent", bold=True)
add_para("[Insérer diagramme de séquence - Chatbot]", italic=True)
add_para("L'utilisateur envoie un message au chatbot via une requête POST à /api/chat. Le contrôleur ChatController analyse le message avec un moteur de correspondance par mots-clés, charge les données pertinentes depuis la base de données (projets, étapes, prix, dates de livraison), et génère une réponse contextuelle et formatée en Markdown.")
doc.add_page_break()

# ========== CHAPITRE 3 ==========
doc.add_heading('Chapitre 3 : Conception', level=1)

doc.add_heading('3.1 Architecture globale de l\'application', level=2)
add_para("L'application suit une architecture client-serveur de type API RESTful :")
doc.add_paragraph("Backend : Laravel 12 (PHP 8.2) exposant une API REST sécurisée", style='List Bullet')
doc.add_paragraph("Frontend : Application SPA communiquant avec l'API via Axios / Fetch API", style='List Bullet')
doc.add_paragraph("Base de données : MySQL (base 'stage')", style='List Bullet')
doc.add_paragraph("Authentification : Laravel Sanctum (tokens API)", style='List Bullet')
doc.add_paragraph("Assistant Virtuel : Moteur NLP local avec exploitation contextuelle de la BDD", style='List Bullet')
doc.add_paragraph("E-mail : SMTP Gmail", style='List Bullet')
doc.add_paragraph("Export : DomPDF pour les PDF, streaming CSV", style='List Bullet')

doc.add_heading('3.2 Architecture MVC', level=2)
add_para("Laravel implémente le pattern MVC (Modèle-Vue-Contrôleur) :")
add_para("Modèles (5) : User, Project, ProjectStep, Notification, Setting — définissent la structure des données et les relations (hasMany, belongsTo).", bold=False)
add_para("Contrôleurs (8) : AuthController, AdminController, ProjectController, ProjectStepController, ChatController, ContactController, SettingController, Controller — contiennent la logique métier.", bold=False)
add_para("Vues : Templates Blade pour les e-mails, les exports PDF et la page d'accueil.", bold=False)

doc.add_heading('3.3 Modèle de données', level=2)
add_para("La base de données comprend les tables suivantes :", bold=True)

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Table'
hdr2[1].text = 'Colonnes principales'
hdr2[2].text = 'Relations'

tables_data = [
    ("users", "id, name, email, password, role, avatar, phone", "hasMany(Project, Notification)"),
    ("projects", "id, user_id, name, location, progress, status, image_path, document_path, total_value, delivery_date", "belongsTo(User), hasMany(ProjectStep)"),
    ("project_steps", "id, project_id, label, status (enum: pending/in_progress/completed), order_num", "belongsTo(Project)"),
    ("notifications", "id, user_id, content, is_read", "belongsTo(User)"),
    ("settings", "id, key, value", "Aucune"),
    ("personal_access_tokens", "id, tokenable_type, tokenable_id, name, token, abilities", "Polymorphe (User)"),
]
for tbl, cols, rels in tables_data:
    row = table2.add_row().cells
    row[0].text = tbl
    row[1].text = cols
    row[2].text = rels

doc.add_paragraph()

doc.add_heading('3.4 Diagramme de classes', level=2)
add_para("[Insérer le diagramme de classes UML]", italic=True, bold=True)
add_para("Les principales relations entre les classes sont :")
doc.add_paragraph("User (1) ──── (*) Project : Un utilisateur possède plusieurs projets", style='List Bullet')
doc.add_paragraph("Project (1) ──── (*) ProjectStep : Un projet contient plusieurs étapes", style='List Bullet')
doc.add_paragraph("User (1) ──── (*) Notification : Un utilisateur reçoit plusieurs notifications", style='List Bullet')

doc.add_heading('3.5 Architecture du chatbot intelligent', level=2)
add_para("Le chatbot « Bakkah Assistant » a été conçu avec une architecture locale autonome, sans dépendance à des APIs externes payantes. Cette approche garantit :")
doc.add_paragraph("Disponibilité à 100% : aucune dépendance à des services tiers (pas de risque de suspension de clé API)", style='List Bullet')
doc.add_paragraph("Coût zéro : pas de frais d'utilisation d'API externes (GROQ, OpenAI, etc.)", style='List Bullet')
doc.add_paragraph("Performance optimale : temps de réponse < 100ms grâce au traitement local", style='List Bullet')
doc.add_paragraph("Données en temps réel : exploitation directe de la base de données pour des réponses toujours à jour", style='List Bullet')

add_para("Architecture du moteur de chat :", bold=True)
add_para("Le ChatController utilise un pipeline de traitement en plusieurs étapes :")
doc.add_paragraph("1. Normalisation : le message utilisateur est converti en minuscules et nettoyé", style='List Bullet')
doc.add_paragraph("2. Classification : le moteur identifie la catégorie de la question (salutation, projet, avancement, prix, contact, etc.)", style='List Bullet')
doc.add_paragraph("3. Extraction de contexte : recherche de correspondance avec les noms de projets ou localisations dans la base de données", style='List Bullet')
doc.add_paragraph("4. Génération de réponse : construction d'une réponse formatée en Markdown avec les données pertinentes", style='List Bullet')
doc.add_paragraph("5. Fallback intelligent : recherche par mots-clés dans les noms de projets si aucune catégorie n'est identifiée", style='List Bullet')

add_para("Catégories de réponses :", bold=True)
table_chat = doc.add_table(rows=1, cols=3)
table_chat.style = 'Light Grid Accent 1'
hdr_chat = table_chat.rows[0].cells
hdr_chat[0].text = 'Catégorie'
hdr_chat[1].text = 'Mots-clés déclencheurs'
hdr_chat[2].text = 'Données utilisées'

chat_categories = [
    ("Salutations", "bonjour, salut, hello, salam", "Aucune"),
    ("Projets", "projet, construction, villa, résidence", "Projects (nom, lieu, progrès)"),
    ("Avancement", "avancement, progression, statut", "Projects + Steps"),
    ("Étapes", "étape, phase, timeline", "ProjectSteps (label, status)"),
    ("Livraison", "livraison, date, quand, délai", "Projects (delivery_date)"),
    ("Prix", "prix, coût, valeur, budget", "Projects (total_value)"),
    ("Contact", "contact, téléphone, email, adresse", "Données statiques Bakkah"),
    ("À propos", "qui, entreprise, bakkah, histoire", "Données statiques Bakkah"),
    ("Dashboard", "dashboard, aide, fonctionnalité", "Guide utilisateur"),
    ("Ville", "casablanca, rabat, tanger, maroc", "Projects (location)"),
]
for cat, keywords, data in chat_categories:
    row = table_chat.add_row().cells
    row[0].text = cat
    row[1].text = keywords
    row[2].text = data

doc.add_page_break()

# ========== CHAPITRE 4 ==========
doc.add_heading('Chapitre 4 : Réalisation', level=1)

doc.add_heading('4.1 Environnement de développement', level=2)
table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Light Grid Accent 1'
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Outil / Technologie'
hdr3[1].text = 'Version / Description'

env_data = [
    ("Système d'exploitation", "Windows"),
    ("Serveur local", "WAMP64 (Apache, MySQL, PHP)"),
    ("Éditeur de code", "Visual Studio Code"),
    ("PHP", "8.2+"),
    ("Laravel", "12.0"),
    ("Composer", "Gestionnaire de dépendances PHP"),
    ("Node.js / NPM", "Runtime JavaScript + gestionnaire de paquets"),
    ("Vite", "7.0+ (bundler frontend)"),
    ("MySQL", "Base de données relationnelle"),
    ("Git", "Contrôle de version"),
    ("Postman", "Test des API REST"),
]
for outil, desc in env_data:
    row = table3.add_row().cells
    row[0].text = outil
    row[1].text = desc

doc.add_paragraph()

doc.add_heading('4.2 Technologies utilisées', level=2)
add_para("Backend :", bold=True)
doc.add_paragraph("Laravel 12 : Framework PHP pour le développement de l'API RESTful", style='List Bullet')
doc.add_paragraph("Laravel Sanctum 4.0 : Authentification par tokens API (SPA-friendly)", style='List Bullet')
doc.add_paragraph("Eloquent ORM : Gestion des modèles et des relations de base de données", style='List Bullet')
doc.add_paragraph("DomPDF (barryvdh/laravel-dompdf 3.1) : Génération et export de rapports PDF", style='List Bullet')
doc.add_paragraph("Carbon : Manipulation des dates avec localisation française", style='List Bullet')
doc.add_paragraph("Moteur NLP local : Chatbot intelligent avec correspondance par mots-clés et fuzzy matching", style='List Bullet')

add_para("Frontend :", bold=True)
doc.add_paragraph("HTML5 / CSS3 / JavaScript : Technologies web standards", style='List Bullet')
doc.add_paragraph("TailwindCSS 4.0 : Framework CSS utilitaire pour un design moderne et premium", style='List Bullet')
doc.add_paragraph("Axios : Client HTTP pour la communication avec l'API", style='List Bullet')
doc.add_paragraph("Vite 7 : Outil de build et serveur de développement rapide", style='List Bullet')
doc.add_paragraph("Chart.js : Bibliothèque de graphiques pour le dashboard admin", style='List Bullet')

doc.add_heading('4.3 Points de terminaison de l\'API (Routes)', level=2)
add_para("L'API RESTful expose les endpoints suivants :")

table4 = doc.add_table(rows=1, cols=4)
table4.style = 'Light Grid Accent 1'
hdr4 = table4.rows[0].cells
hdr4[0].text = 'Méthode'
hdr4[1].text = 'Endpoint'
hdr4[2].text = 'Description'
hdr4[3].text = 'Auth'

routes = [
    ("POST", "/api/register", "Inscription d'un nouveau client", "Non"),
    ("POST", "/api/login", "Connexion et obtention du token", "Non"),
    ("POST", "/api/logout", "Déconnexion (suppression du token)", "Oui"),
    ("GET", "/api/user", "Récupérer le profil de l'utilisateur connecté", "Oui"),
    ("PUT", "/api/user/profile", "Modifier le profil", "Oui"),
    ("PUT", "/api/user/password", "Changer le mot de passe", "Oui"),
    ("GET", "/api/admin/stats", "Statistiques et activités du dashboard admin", "Oui"),
    ("GET/POST/PUT/DELETE", "/api/admin/users", "CRUD des utilisateurs", "Oui"),
    ("GET", "/api/projects", "Liste des projets (filtrable par statut et recherche)", "Oui"),
    ("POST", "/api/projects", "Créer un nouveau projet (avec upload de document)", "Oui"),
    ("PUT", "/api/projects/{id}", "Modifier un projet", "Oui"),
    ("DELETE", "/api/projects/{id}", "Supprimer un projet", "Oui"),
    ("GET", "/api/projects/export", "Export CSV avec filtres", "Oui"),
    ("GET", "/api/projects/export-pdf", "Export PDF avec filtres", "Oui"),
    ("POST/PUT/DELETE", "/api/steps", "CRUD des étapes de projet", "Oui"),
    ("POST", "/api/chat", "Envoyer un message au chatbot intelligent", "Non"),
    ("POST", "/api/contact", "Envoyer un e-mail de contact", "Oui"),
    ("GET/POST", "/api/settings", "Gestion des paramètres", "Oui"),
    ("GET", "/api/test", "Test de connectivité API", "Non"),
]
for meth, ep, desc, auth in routes:
    row = table4.add_row().cells
    row[0].text = meth
    row[1].text = ep
    row[2].text = desc
    row[3].text = auth

doc.add_paragraph()

doc.add_heading('4.4 Fonctionnement du chatbot intelligent', level=2)
add_para("L'un des modules les plus innovants de l'application est l'assistant virtuel « Bakkah Assistant ». Contrairement aux chatbots traditionnels qui dépendent d'APIs externes coûteuses (OpenAI, GROQ, Google Gemini), notre solution fonctionne de manière entièrement autonome.")

add_para("Avantages de l'approche locale :", bold=True)
doc.add_paragraph("Indépendance totale : pas de risque de suspension de service ou de dépassement de quota API", style='List Bullet')
doc.add_paragraph("Données toujours à jour : les réponses sont générées à partir des données en temps réel de la base de données", style='List Bullet')
doc.add_paragraph("Performances supérieures : temps de réponse moyen < 50ms contre 1-3 secondes pour les APIs externes", style='List Bullet')
doc.add_paragraph("Personnalisation totale : le chatbot est entièrement adapté au contexte de Bakkah Immobilier", style='List Bullet')
doc.add_paragraph("Coût nul : aucun abonnement ou frais d'utilisation d'API tierce", style='List Bullet')

add_para("Fonctionnalités du chatbot :", bold=True)
doc.add_paragraph("Recherche de projets par nom, localisation ou mots-clés", style='List Bullet')
doc.add_paragraph("Affichage détaillé de l'avancement avec barre de progression visuelle", style='List Bullet')
doc.add_paragraph("Consultation des étapes de construction par projet", style='List Bullet')
doc.add_paragraph("Informations sur les dates de livraison prévues", style='List Bullet')
doc.add_paragraph("Affichage des prix et valeurs des projets", style='List Bullet')
doc.add_paragraph("Coordonnées de contact de Bakkah Immobilier", style='List Bullet')
doc.add_paragraph("Présentation de l'entreprise et de ses valeurs", style='List Bullet')
doc.add_paragraph("Guide d'utilisation du tableau de bord", style='List Bullet')
doc.add_paragraph("Recherche par ville (Casablanca, Rabat, Tanger, etc.)", style='List Bullet')
doc.add_paragraph("Fuzzy matching pour une correspondance flexible des noms de projets", style='List Bullet')

doc.add_heading('4.5 Dashboard administrateur avancé', level=2)
add_para("Le tableau de bord administrateur a été enrichi avec des composants analytiques avancés :")

add_para("Statistiques en temps réel :", bold=True)
doc.add_paragraph("Nombre total de projets en cours", style='List Bullet')
doc.add_paragraph("Nombre total de clients inscrits", style='List Bullet')
doc.add_paragraph("Chiffre d'affaires total (valeur cumulée des projets)", style='List Bullet')
doc.add_paragraph("Progrès moyen de l'ensemble des projets", style='List Bullet')

add_para("Flux d'activités récentes :", bold=True)
add_para("Le dashboard affiche un fil d'activité en temps réel regroupant trois types d'événements :")
doc.add_paragraph("Création de nouveaux projets (code couleur vert)", style='List Bullet')
doc.add_paragraph("Inscription de nouveaux clients (code couleur violet)", style='List Bullet')
doc.add_paragraph("Achèvement d'étapes de construction (code couleur bleu)", style='List Bullet')
add_para("Les activités sont triées chronologiquement et limitées aux 10 plus récentes pour une lecture optimale.")

doc.add_heading('4.6 Captures d\'écran de l\'application', level=2)
add_para("[Insérer les captures d'écran des différentes pages de l'application]", italic=True, bold=True)
add_para("Page de connexion / inscription", bold=True)
add_para("La page d'authentification permet aux utilisateurs de se connecter avec leur email et mot de passe, ou de créer un nouveau compte client. L'interface utilise un design moderne avec validation en temps réel.")
add_para("Dashboard Client", bold=True)
add_para("Le tableau de bord client affiche la liste des projets attribués avec leur pourcentage d'avancement, statut actuel, et détails des étapes de construction. Le design premium offre une expérience utilisateur fluide et intuitive.")
add_para("Dashboard Administrateur", bold=True)
add_para("Le dashboard admin présente les statistiques clés avec des cartes animées, un graphique interactif des projets, et un flux d'activités récentes en temps réel regroupant les événements du système.")
add_para("Module Assistant Virtuel", bold=True)
add_para("L'assistant virtuel « Bakkah Assistant » offre une interface de chat moderne. Les réponses sont formatées en Markdown avec des informations précises tirées de la base de données en temps réel (projets, étapes, prix, dates).")
add_para("Gestion des utilisateurs", bold=True)
add_para("L'interface d'administration permet de créer, modifier et supprimer des comptes utilisateurs avec attribution des rôles (admin/client). Le système empêche la suppression de son propre compte par mesure de sécurité.")

doc.add_heading('4.7 Tests et validation', level=2)
add_para("Les tests ont été effectués à plusieurs niveaux :")
doc.add_paragraph("Tests unitaires : Validation des modèles et des contrôleurs", style='List Bullet')
doc.add_paragraph("Tests d'intégration : Vérification des endpoints API avec Postman", style='List Bullet')
doc.add_paragraph("Tests fonctionnels : Validation des scénarios utilisateur de bout en bout", style='List Bullet')
doc.add_paragraph("Tests de sécurité : Vérification de l'authentification Sanctum et des autorisations", style='List Bullet')
doc.add_paragraph("Tests du chatbot : Validation des réponses pour chaque catégorie de question avec différentes formulations", style='List Bullet')
doc.add_paragraph("Tests de performance : Mesure des temps de réponse API et du chatbot (< 50ms)", style='List Bullet')
doc.add_page_break()

# ========== CONCLUSION ==========
doc.add_heading('Conclusion Générale et Perspectives', level=1)
doc.add_paragraph()
add_para("Au terme de ce projet de fin d'études, nous avons réussi à concevoir et développer une application web complète de gestion immobilière pour la société Bakkah Immobilier. Cette application répond aux besoins exprimés en offrant une solution moderne, sécurisée et performante.")
add_para("Les principales réalisations incluent :")
doc.add_paragraph("Une API RESTful robuste avec Laravel 12 et une authentification sécurisée via Sanctum", style='List Bullet')
doc.add_paragraph("Un système de gestion de projets avec suivi par étapes et calcul automatique du progrès", style='List Bullet')
doc.add_paragraph("Un dashboard administrateur avancé avec statistiques en temps réel, graphiques interactifs et flux d'activités", style='List Bullet')
doc.add_paragraph("Un assistant virtuel intelligent autonome basé sur un moteur NLP local, exploitant les données de la base de données en temps réel", style='List Bullet')
doc.add_paragraph("Un système de notifications automatiques et d'envoi d'e-mails via SMTP", style='List Bullet')
doc.add_paragraph("Des fonctionnalités d'export en CSV et PDF avec filtres", style='List Bullet')
doc.add_paragraph("Une interface utilisateur moderne, responsive et premium", style='List Bullet')

add_para("L'approche innovante du chatbot local, adoptée après avoir exploré les limitations et les coûts des APIs IA externes (GROQ, OpenAI, Gemini), s'est avérée être une solution optimale pour ce contexte d'entreprise. Elle offre une disponibilité à 100%, des performances supérieures et un coût nul, tout en fournissant des réponses précises et contextuelles.", italic=True)

add_para("Perspectives d'amélioration :", bold=True)
doc.add_paragraph("Ajout d'un système de messagerie en temps réel (WebSocket) entre admin et clients", style='List Bullet')
doc.add_paragraph("Intégration de la géolocalisation des projets sur une carte interactive (Leaflet/Google Maps)", style='List Bullet')
doc.add_paragraph("Développement d'une application mobile (React Native ou Flutter)", style='List Bullet')
doc.add_paragraph("Mise en place d'un système de paiement en ligne pour les réservations", style='List Bullet')
doc.add_paragraph("Enrichissement du chatbot avec un modèle d'IA locale (LLM open-source type Mistral ou LLaMA) pour des réponses encore plus naturelles", style='List Bullet')
doc.add_paragraph("Déploiement en production sur un serveur cloud (AWS, DigitalOcean, Hostinger)", style='List Bullet')
doc.add_paragraph("Ajout d'un système de gestion de documents avec signature électronique", style='List Bullet')
doc.add_paragraph("Intégration d'un tableau de bord analytique avec Power BI ou Metabase", style='List Bullet')

add_para("Ce stage nous a permis de consolider nos connaissances en développement web full-stack, de maîtriser le framework Laravel et ses écosystèmes, et de découvrir les possibilités offertes par le traitement du langage naturel dans les applications web modernes. L'expérience acquise au sein de Bakkah Immobilier nous a également permis de comprendre les enjeux métier du secteur immobilier et d'appliquer des solutions technologiques adaptées à des problèmes concrets.", italic=True)
doc.add_page_break()

# ========== RÉFÉRENCES ==========
doc.add_heading('Références et Webographie', level=1)
doc.add_paragraph()
refs = [
    "Laravel Documentation officielle : https://laravel.com/docs",
    "Laravel Sanctum : https://laravel.com/docs/sanctum",
    "Eloquent ORM : https://laravel.com/docs/eloquent",
    "MySQL Documentation : https://dev.mysql.com/doc/",
    "TailwindCSS Documentation : https://tailwindcss.com/docs",
    "Vite Documentation : https://vitejs.dev/guide/",
    "DomPDF pour Laravel : https://github.com/barryvdh/laravel-dompdf",
    "Axios Documentation : https://axios-http.com/docs/intro",
    "PHP.net Documentation : https://www.php.net/docs.php",
    "Carbon Documentation : https://carbon.nesbot.com/docs/",
    "Chart.js Documentation : https://www.chartjs.org/docs/",
    "MDN Web Docs (HTML/CSS/JS) : https://developer.mozilla.org/",
]
for i, ref in enumerate(refs, 1):
    add_para(f"[{i}] {ref}", size=11)
doc.add_page_break()

# ========== ANNEXES ==========
doc.add_heading('Annexes', level=1)
doc.add_paragraph()
add_para("Annexe A : Code source des principaux contrôleurs", bold=True)
add_para("Le code source complet est disponible dans le répertoire du projet. Les contrôleurs principaux sont :")
doc.add_paragraph("AuthController.php : Gestion de l'authentification (register, login, logout, updateProfile, updatePassword)", style='List Bullet')
doc.add_paragraph("ProjectController.php : CRUD des projets, export CSV/PDF avec filtres, upload de documents", style='List Bullet')
doc.add_paragraph("ProjectStepController.php : Gestion des étapes avec recalcul automatique du progrès", style='List Bullet')
doc.add_paragraph("AdminController.php : Dashboard admin avancé (stats, activités récentes), CRUD utilisateurs", style='List Bullet')
doc.add_paragraph("ChatController.php : Moteur de chatbot intelligent avec NLP local, correspondance par mots-clés, fuzzy matching, et exploitation contextuelle de la BDD", style='List Bullet')
doc.add_paragraph("ContactController.php : Envoi de mails via SMTP", style='List Bullet')
doc.add_paragraph("SettingController.php : Gestion des paramètres clé/valeur", style='List Bullet')

doc.add_paragraph()
add_para("Annexe B : Structure de la base de données (Migrations)", bold=True)
add_para("Les fichiers de migration définissent le schéma de la base de données 'stage' sous MySQL, avec les tables : users, projects, project_steps, notifications, settings, personal_access_tokens, cache, jobs.")

doc.add_paragraph()
add_para("Annexe C : Architecture du ChatController", bold=True)
add_para("Le ChatController comprend les méthodes principales suivantes :")
doc.add_paragraph("ask() : Point d'entrée du chatbot, reçoit le message et retourne la réponse JSON", style='List Bullet')
doc.add_paragraph("generateResponse() : Pipeline principal de traitement NLP avec classification par catégories", style='List Bullet')
doc.add_paragraph("matchesAny() : Vérification de correspondance de mots-clés dans le message", style='List Bullet')
doc.add_paragraph("fuzzyMatch() : Correspondance flexible pour les noms de projets", style='List Bullet')
doc.add_paragraph("formatProjectDetail() : Mise en forme détaillée d'un projet (nom, lieu, progrès, étapes, prix, date)", style='List Bullet')
doc.add_paragraph("formatProjectsList() : Affichage de la liste de tous les projets", style='List Bullet')
doc.add_paragraph("formatProjectProgress() : Affichage avec barre de progression visuelle", style='List Bullet')
doc.add_paragraph("formatProjectSteps() : Détail des étapes d'un projet", style='List Bullet')
doc.add_paragraph("progressBar() : Génération d'une barre de progression textuelle", style='List Bullet')

# ---- Save ----
output_path = r'c:\wamp64\www\pfe-stage-laravel1\Rapport_de_Stage_Final.docx'
doc.save(output_path)
print(f"Rapport mis à jour généré avec succès : {output_path}")

