from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = r'c:\wamp64\www\pfe-stage-laravel1\Rapport_de_Stage_Bakkah_Immobilier1 .docx'

try:
    doc = Document(doc_path)
    
    # Access the first section's header
    section = doc.sections[0]
    header = section.header
    
    # Clear existing header content
    for paragraph in header.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None
    
    for table in header.tables:
        t = table._element
        t.getparent().remove(t)
        table._tbl = table._element = None
        
    # Re-add an empty paragraph required at the top if deleted
    if not header.paragraphs:
        header.add_paragraph()
        
    header_table = header.add_table(rows=1, cols=3, width=Inches(6.5))
    header_table.autofit = False
    header_table.columns[0].width = Inches(2.0)
    header_table.columns[1].width = Inches(2.5)
    header_table.columns[2].width = Inches(2.0)

    # OFPPT Logo (Left)
    cell_left = header_table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    run_left = p_left.add_run()
    try:
        run_left.add_picture(r"c:\wamp64\www\pfe-stage-laravel1\ofppt.jpg", width=Inches(1.2))
    except Exception as e:
        print(f"Erreur chargement logo OFPPT: {e}")
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Bakkah Logo (Right)
    cell_right = header_table.cell(0, 2)
    p_right = cell_right.paragraphs[0]
    run_right = p_right.add_run()
    try:
        run_right.add_picture(r"c:\wamp64\www\pfe-stage-laravel1\bakkah.png", width=Inches(1.2))
    except Exception as e:
        print(f"Erreur chargement logo Bakkah: {e}")
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(doc_path)
    print("Succès")
except PermissionError:
    print("Erreur : Le fichier est ouvert dans Word. Veuillez le fermer.")
except Exception as e:
    print(f"Erreur inattendue : {e}")
